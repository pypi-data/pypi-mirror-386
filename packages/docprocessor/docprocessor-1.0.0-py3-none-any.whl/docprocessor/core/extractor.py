# backend/app/services/search/content_extractor.py

"""
Content extraction service for semantic search.

Extracts text content from various file formats for indexing.
"""

import logging
from pathlib import Path
from typing import Any, Dict

from .ocr import extract_pdf_for_llm

logger = logging.getLogger(__name__)


class ContentExtractionError(Exception):
    """Raised when content extraction fails."""

    pass


class ContentExtractor:
    """
    Extracts text content from various file formats.

    Supported formats:
    - PDF: Using OCR pipeline (extract_pdf_for_llm)
    - TXT/MD: Direct text read
    - DOCX: python-docx extraction
    - Images: OCR fallback
    """

    def __init__(self):
        self.supported_extensions = {
            ".pdf",
            ".txt",
            ".md",
            ".docx",
            ".png",
            ".jpg",
            ".jpeg",
            ".gif",
            ".bmp",
        }

    def extract(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text content from a file.

        Args:
            file_path: Path to the file to extract

        Returns:
            Dictionary with:
                - text: Extracted text content
                - page_count: Number of pages (for PDFs)
                - metadata: Additional metadata

        Raises:
            ContentExtractionError: If extraction fails
        """
        if not file_path.exists():
            raise ContentExtractionError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension not in self.supported_extensions:
            logger.warning(f"Unsupported file type: {extension}, attempting text read")
            return self._extract_as_text(file_path)

        try:
            if extension == ".pdf":
                return self._extract_pdf(file_path)
            elif extension in {".txt", ".md"}:
                return self._extract_text(file_path)
            elif extension == ".docx":
                return self._extract_docx(file_path)
            elif extension in {".png", ".jpg", ".jpeg", ".gif", ".bmp"}:
                return self._extract_image(file_path)
            else:
                return self._extract_as_text(file_path)
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            raise ContentExtractionError(f"Failed to extract content: {str(e)}")

    def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF using OCR pipeline."""
        logger.info(f"Extracting PDF: {file_path}")

        with open(file_path, "rb") as f:
            pdf_bytes = f.read()

        text = extract_pdf_for_llm(pdf_bytes)

        # Count pages from text markers
        page_count = text.count("<page_")

        return {
            "text": text,
            "page_count": page_count if page_count > 0 else 1,
            "metadata": {"format": "pdf", "extraction_method": "ocr_pipeline"},
        }

    def _extract_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files."""
        logger.info(f"Reading text file: {file_path}")

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()

        return {
            "text": text,
            "page_count": 1,
            "metadata": {"format": file_path.suffix[1:], "extraction_method": "direct_read"},
        }

    def _extract_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        logger.info(f"Extracting DOCX: {file_path}")

        try:
            import docx
        except ImportError:
            raise ContentExtractionError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        try:
            doc = docx.Document(str(file_path))

            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)

            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                tables_text.append("\n".join(table_data))

            if tables_text:
                text += "\n\n" + "\n\n".join(tables_text)

            return {
                "text": text,
                "page_count": 1,  # DOCX doesn't have clear page boundaries
                "metadata": {
                    "format": "docx",
                    "extraction_method": "python-docx",
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                },
            }
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise ContentExtractionError(f"Failed to extract DOCX: {str(e)}")

    def _extract_image(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from images using OCR."""
        logger.info(f"Extracting image with OCR: {file_path}")

        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            raise ContentExtractionError("PIL and pytesseract required for image extraction")

        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)

            return {
                "text": text,
                "page_count": 1,
                "metadata": {
                    "format": file_path.suffix[1:],
                    "extraction_method": "tesseract_ocr",
                    "image_size": image.size,
                },
            }
        except Exception as e:
            logger.error(f"Error extracting image: {e}")
            raise ContentExtractionError(f"Failed to extract image: {str(e)}")

    def _extract_as_text(self, file_path: Path) -> Dict[str, Any]:
        """Fallback: attempt to read as text."""
        logger.warning(f"Attempting fallback text extraction for: {file_path}")

        try:
            return self._extract_text(file_path)
        except Exception as e:
            raise ContentExtractionError(
                f"Cannot extract content from {file_path.suffix} file: {str(e)}"
            )

    def is_supported(self, extension: str) -> bool:
        """Check if a file extension is supported."""
        return extension.lower() in self.supported_extensions


# Global instance
content_extractor = ContentExtractor()
