#!/usr/bin/env python3
"""
Universal XMind Converter

A Python tool that converts multiple file formats (TXT, HTML, Word, Excel, Markdown) to XMind mind map format.
Supports automatic format detection and extensible parser architecture.

Version: 2.0
"""

import json
import zipfile
import re
import os
from datetime import datetime
from pathlib import Path
import mimetypes

# 可选依赖，用于处理Word和Excel文件
try:
    import docx
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False


def escape_xml_text(text):
    """Escape XML special characters"""
    if not text:
        return ""
    return (text.replace("&", "&amp;")
               .replace("<", "&lt;")
               .replace(">", "&gt;")
               .replace("\"", "&quot;")
               .replace("'", "&apos;"))


def generate_id():
    """Generate unique ID"""
    import uuid
    return str(uuid.uuid4()).replace('-', '')


def create_json_structure(title, children):
    """Create JSON structure"""
    return {
        "id": generate_id(),
        "class": "sheet",
        "rootTopic": {
            "id": generate_id(),
            "class": "topic",
            "title": title,
            "structureClass": "org.xmind.ui.logic.right",
            "children": {
                "attached": children
            }
        },
        "title": "画布 1",
        "extensions": [{
            "provider": "org.xmind.ui.skeleton.structure.style",
            "content": {
                "centralTopic": "org.xmind.ui.logic.right"
            }
        }],
        "theme": {
            "map": {
                "id": generate_id(),
                "properties": {
                    "svg:fill": "#ffffff",
                    "multi-line-colors": "#F9423A #F6A04D #F3D321 #00BC7B #486AFF #4D49BE",
                    "color-list": "#000229 #1F2766 #52CC83 #4D86DB #99142F #245570",
                    "line-tapered": "none"
                }
            },
            "centralTopic": {
                "id": generate_id(),
                "properties": {
                    "fo:font-family": "Droid Serif",
                    "fo:font-size": "30pt",
                    "fo:font-weight": "400",
                    "fo:font-style": "normal",
                    "fo:color": "inherited",
                    "fo:text-transform": "manual",
                    "fo:text-decoration": "none",
                    "fo:text-align": "center",
                    "svg:fill": "#000229",
                    "fill-pattern": "solid",
                    "line-width": "3pt",
                    "line-color": "#000229",
                    "line-pattern": "solid",
                    "border-line-color": "inherited",
                    "border-line-width": "3pt",
                    "border-line-pattern": "inherited",
                    "shape-class": "org.xmind.topicShape.roundedRect",
                    "line-class": "org.xmind.branchConnection.roundedElbow",
                    "arrow-end-class": "org.xmind.arrowShape.none",
                    "alignment-by-level": "inactived"
                }
            },
            "mainTopic": {
                "id": generate_id(),
                "properties": {
                    "fo:font-family": "Droid Serif",
                    "fo:font-size": "18pt",
                    "fo:font-weight": "400",
                    "fo:font-style": "normal",
                    "fo:color": "inherited",
                    "fo:text-transform": "manual",
                    "fo:text-decoration": "none",
                    "fo:text-align": "left",
                    "svg:fill": "inherited",
                    "fill-pattern": "none",
                    "line-width": "inherited",
                    "line-color": "inherited",
                    "line-pattern": "inherited",
                    "border-line-color": "inherited",
                    "border-line-width": "0pt",
                    "border-line-pattern": "inherited",
                    "shape-class": "org.xmind.topicShape.roundedRect",
                    "line-class": "org.xmind.branchConnection.roundedElbow",
                    "arrow-end-class": "inherited"
                }
            },
            "subTopic": {
                "id": generate_id(),
                "properties": {
                    "fo:font-family": "Droid Serif",
                    "fo:font-size": "14pt",
                    "fo:font-weight": "400",
                    "fo:font-style": "normal",
                    "fo:color": "inherited",
                    "fo:text-transform": "manual",
                    "fo:text-decoration": "none",
                    "fo:text-align": "left",
                    "svg:fill": "inherited",
                    "fill-pattern": "none",
                    "line-width": "2pt",
                    "line-color": "inherited",
                    "line-pattern": "inherited",
                    "border-line-color": "inherited",
                    "border-line-width": "0pt",
                    "border-line-pattern": "inherited",
                    "shape-class": "org.xmind.topicShape.roundedRect",
                    "line-class": "org.xmind.branchConnection.roundedElbow",
                    "arrow-end-class": "inherited"
                }
            },
            "skeletonThemeId": "c1fbada1b45ba2e3bfc3b8b57b",
            "colorThemeId": "Rainbow-#000229-MULTI_LINE_COLORS"
        }
    }


def create_topic(title, children=None):
    """Create topic node"""
    topic = {
        "id": generate_id(),
        "title": title
    }
    
    if children:
        topic["children"] = {
            "attached": children
        }
    
    return topic


# ==============================================
# 文件格式解析器基类
# ==============================================

class BaseParser:
    """基础解析器类"""
    
    def __init__(self, file_path):
        self.file_path = file_path
    
    def parse(self):
        """解析文件并返回JSON结构"""
        raise NotImplementedError("子类必须实现parse方法")
    
    def extract_title(self, content):
        """提取标题"""
        return Path(self.file_path).stem


# ==============================================
# Markdown解析器
# ==============================================

class MarkdownParser(BaseParser):
    """Markdown文件解析器"""
    
    def parse(self):
        """解析Markdown文件"""
        with open(self.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        title = self.extract_title_from_content(lines)
        
        # 构建层级结构
        topics_by_level = {}
        root_children = []
        title_extracted = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检测标题级别
            if line.startswith('#### '):
                level = 4
                title_text = line[5:].strip()
            elif line.startswith('### '):
                level = 3
                title_text = line[4:].strip()
            elif line.startswith('## '):
                level = 2
                title_text = line[3:].strip()
            elif line.startswith('# '):
                level = 1
                title_text = line[2:].strip()
                # 第一个一级标题作为文档标题，不处理为节点
                if not title_extracted:
                    title_extracted = True
                    continue
            else:
                continue
            
            # 创建topic
            topic = create_topic(title_text)
            topics_by_level[level] = topic
            
            # 添加到父节点
            if level == 1:
                # 一级标题直接作为根节点的子节点
                root_children.append(topic)
            else:
                # 其他级别的标题，找到父节点（级别-1）
                parent_level = level - 1
                parent_found = False
                
                # 从当前级别往上找，直到找到合适的父节点
                while parent_level >= 1 and not parent_found:
                    if parent_level in topics_by_level:
                        parent_topic = topics_by_level[parent_level]
                        if "children" not in parent_topic:
                            parent_topic["children"] = {"attached": []}
                        parent_topic["children"]["attached"].append(topic)
                        parent_found = True
                    else:
                        parent_level -= 1
                
                # 如果没找到父节点，就作为根节点的子节点
                if not parent_found:
                    root_children.append(topic)
        
        return create_json_structure(title, root_children)
    
    def extract_title_from_content(self, lines):
        """从内容中提取标题"""
        for line in lines:
            if line.startswith('# '):
                return line[2:].strip()
        return Path(self.file_path).stem
    
    def parse_markdown_to_json(self, markdown_file):
        """解析Markdown文件并转换为JSON结构（兼容旧版本）"""
        return self.parse()


# ==============================================
# 文本大纲解析器
# ==============================================

class TextOutlineParser(BaseParser):
    """文本大纲解析器 - 支持缩进层级"""
    
    def parse(self):
        """解析文本大纲文件"""
        with open(self.file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        title = self.extract_title(lines)
        
        # 构建层级结构
        topics_by_level = {}
        root_children = []
        
        # 跳过标题行（如果标题行存在）
        title_found = False
        for line in lines:
            line = line.rstrip('\n\r')
            if not line.strip():
                continue
            
            # 检查是否是标题行
            stripped_line = line.lstrip()
            if stripped_line.strip() == title.strip() and not title_found:
                title_found = True
                continue  # 跳过标题行
            
            # 计算缩进级别
            indent_level = len(line) - len(stripped_line)
            level = self.indent_to_level(indent_level)
            title_text = stripped_line.strip('-*• ')
            
            if not title_text:
                continue
            
            # 创建topic
            topic = create_topic(title_text)
            topics_by_level[level] = topic
            
            # 添加到父节点
            if level == 0:
                root_children.append(topic)
            else:
                parent_level = level - 1
                if parent_level in topics_by_level:
                    parent_topic = topics_by_level[parent_level]
                    if "children" not in parent_topic:
                        parent_topic["children"] = {"attached": []}
                    parent_topic["children"]["attached"].append(topic)
                else:
                    # 父层级不存在，直接添加到根节点
                    root_children.append(topic)
        
        return create_json_structure(title, root_children)
    
    def indent_to_level(self, indent):
        """将缩进转换为层级"""
        # 假设每4个空格或1个制表符为一级
        return indent // 4
    
    def extract_title(self, lines):
        """提取标题"""
        for line in lines:
            stripped = line.strip()
            if stripped and not stripped.startswith((' ', '\t', '-', '*', '•')):
                return stripped
        return Path(self.file_path).stem


# ==============================================
# HTML解析器
# ==============================================

class HtmlParser(BaseParser):
    """HTML文件解析器"""
    
    def __init__(self, file_path):
        super().__init__(file_path)
        if not HTML_AVAILABLE:
            raise ImportError("需要安装BeautifulSoup4: pip install beautifulsoup4")
    
    def parse(self):
        """解析HTML文件"""
        with open(self.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        soup = BeautifulSoup(content, 'html.parser')
        title = self.extract_title(soup)
        
        # 查找标题结构 (h1-h6)
        root_children = []
        topics_by_level = {}
        
        for heading_tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            headings = soup.find_all(heading_tag)
            level = int(heading_tag[1]) - 1  # h1=0, h2=1, etc.
            
            for heading in headings:
                title_text = heading.get_text().strip()
                if not title_text:
                    continue
                
                topic = create_topic(title_text)
                topics_by_level[level] = topic
                
                if level == 0:
                    root_children.append(topic)
                else:
                    parent_level = level - 1
                    if parent_level in topics_by_level:
                        parent_topic = topics_by_level[parent_level]
                        if "children" not in parent_topic:
                            parent_topic["children"] = {"attached": []}
                        parent_topic["children"]["attached"].append(topic)
        
        # 如果没有找到标题，尝试列表结构
        if not root_children:
            root_children = self.parse_lists(soup)
        
        return create_json_structure(title, root_children)
    
    def extract_title(self, soup):
        """提取标题"""
        title_tag = soup.find('title')
        if title_tag:
            return title_tag.get_text().strip()
        
        h1_tag = soup.find('h1')
        if h1_tag:
            return h1_tag.get_text().strip()
        
        return Path(self.file_path).stem
    
    def parse_lists(self, soup):
        """解析列表结构"""
        def parse_list_items(items, level=0):
            topics = []
            for item in items:
                text = item.get_text().strip()
                if text:
                    topic = create_topic(text)
                    
                    # 查找子列表
                    sublist = item.find(['ul', 'ol'])
                    if sublist:
                        sub_items = sublist.find_all('li', recursive=False)
                        if sub_items:
                            sub_topics = parse_list_items(sub_items, level + 1)
                            if sub_topics:
                                topic["children"] = {"attached": sub_topics}
                    
                    topics.append(topic)
            return topics
        
        # 查找顶级列表
        root_topics = []
        for list_tag in soup.find_all(['ul', 'ol']):
            # 只处理顶级列表（不在其他列表内的）
            if not list_tag.find_parent(['ul', 'ol']):
                items = list_tag.find_all('li', recursive=False)
                if items:
                    root_topics.extend(parse_list_items(items))
        
        return root_topics


# ==============================================
# Word文档解析器
# ==============================================

class WordParser(BaseParser):
    """Word文档解析器"""
    
    def __init__(self, file_path):
        super().__init__(file_path)
        if not WORD_AVAILABLE:
            raise ImportError("需要安装python-docx: pip install python-docx")
    
    def parse(self):
        """解析Word文档"""
        doc = docx.Document(self.file_path)
        title = self.extract_title(doc)
        
        # 按段落解析
        root_children = []
        topics_by_level = {}
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 根据样式判断层级
            level = self.get_heading_level(paragraph)
            if level is not None:
                topic = create_topic(text)
                topics_by_level[level] = topic
                
                if level == 0:
                    root_children.append(topic)
                else:
                    parent_level = level - 1
                    if parent_level in topics_by_level:
                        parent_topic = topics_by_level[parent_level]
                        if "children" not in parent_topic:
                            parent_topic["children"] = {"attached": []}
                        parent_topic["children"]["attached"].append(topic)
        
        return create_json_structure(title, root_children)
    
    def extract_title(self, doc):
        """提取标题"""
        # 尝试从属性中获取
        if doc.core_properties.title:
            return doc.core_properties.title
        
        # 查找第一个标题
        for paragraph in doc.paragraphs:
            if paragraph.style and 'Heading' in paragraph.style.name:
                return paragraph.text.strip()
        
        # 使用文件名
        return Path(self.file_path).stem
    
    def get_heading_level(self, paragraph):
        """获取段落标题级别"""
        if not paragraph.style:
            return None
        
        style_name = paragraph.style.name
        
        # Word标准标题样式
        if style_name.startswith('Heading '):
            try:
                level = int(style_name.split()[-1]) - 1
                return min(level, 5)  # 限制最大层级
            except ValueError:
                pass
        
        # 自定义标题样式
        heading_patterns = ['标题', 'Heading', 'head', 'title']
        for i, pattern in enumerate(heading_patterns):
            if pattern.lower() in style_name.lower():
                return i
        
        return None


# ==============================================
# Excel解析器
# ==============================================

class ExcelParser(BaseParser):
    """Excel文件解析器"""
    
    def __init__(self, file_path):
        super().__init__(file_path)
        if not EXCEL_AVAILABLE:
            raise ImportError("需要安装openpyxl: pip install openpyxl")
    
    def parse(self):
        """解析Excel文件"""
        wb = openpyxl.load_workbook(self.file_path)
        
        # 使用第一个工作表
        ws = wb.active
        title = self.extract_title(ws)
        
        # 按层级组织数据
        root_children = []
        
        # 假设第一列是层级，第二列是内容
        # 或者按缩进层级解析
        current_topics = {}  # level -> topic
        
        for row in ws.iter_rows(min_row=1, values_only=True):
            if not row or not row[0]:
                continue
            
            # 尝试不同的解析策略
            level, text = self.parse_row(row)
            if text:
                topic = create_topic(text)
                current_topics[level] = topic
                
                if level == 0:
                    root_children.append(topic)
                else:
                    parent_level = level - 1
                    if parent_level in current_topics:
                        parent_topic = current_topics[parent_level]
                        if "children" not in parent_topic:
                            parent_topic["children"] = {"attached": []}
                        parent_topic["children"]["attached"].append(topic)
        
        return create_json_structure(title, root_children)
    
    def extract_title(self, worksheet):
        """提取标题"""
        # 使用工作表名称
        title = worksheet.title
        
        # 或者使用第一个非空单元格
        if not title or title == 'Sheet1':
            for row in worksheet.iter_rows(min_row=1, max_row=1, values_only=True):
                for cell in row:
                    if cell:
                        title = str(cell)
                        break
        
        return title or Path(self.file_path).stem
    
    def parse_row(self, row):
        """解析行数据"""
        # 策略1: 第一列是层级，第二列是内容
        if len(row) >= 2:
            try:
                level = int(row[0]) if isinstance(row[0], (int, float)) else 0
                text = str(row[1]) if row[1] else ""
                return level, text.strip()
            except (ValueError, TypeError):
                pass
        
        # 策略2: 按缩进或特殊字符判断
        text = str(row[0]) if row[0] else ""
        level = 0
        
        # 计算前导空格或特殊字符
        stripped = text.lstrip()
        if stripped != text:
            indent = len(text) - len(stripped)
            level = indent // 2  # 每2个空格一级
            text = stripped.strip('-*•→')
        
        return level, text.strip()


# ==============================================
# 解析器工厂
# ==============================================

class ParserFactory:
    """解析器工厂类"""
    
    PARSERS = {
        '.md': MarkdownParser,
        '.markdown': MarkdownParser,
        '.txt': TextOutlineParser,
        '.text': TextOutlineParser,
        '.html': HtmlParser,
        '.htm': HtmlParser,
        '.docx': WordParser,
        '.xlsx': ExcelParser,
        '.xls': ExcelParser,
    }
    
    @classmethod
    def detect_file_type(cls, file_path):
        """检测文件类型（通过扩展名和内容）"""
        ext = Path(file_path).suffix.lower()
        
        # 如果有扩展名，优先使用扩展名
        if ext:
            return ext
        
        # 没有扩展名时，通过内容检测
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read(1000)  # 读取前1000字符
                
                # HTML检测
                if content.strip().startswith('<!DOCTYPE') or content.strip().startswith('<html'):
                    return '.html'
                
                # Markdown检测
                if any(line.strip().startswith('#') for line in content.split('\n')[:10]):
                    return '.md'
                
                # OPML检测
                if '<opml' in content.lower():
                    return '.opml'
                
                # 文本大纲检测（有缩进结构）
                lines = content.split('\n')[:20]
                indent_chars = ['  ', '\t', '·', '•', '-']
                if any(any(line.startswith(char) for char in indent_chars) for line in lines if line.strip()):
                    return '.txt'
                
                # 默认文本
                return '.txt'
                
        except Exception:
            return '.txt'
    
    @classmethod
    def get_parser(cls, file_path):
        """根据文件路径获取相应的解析器"""
        file_ext = cls.detect_file_type(file_path)
        
        if file_ext not in cls.PARSERS:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        parser_class = cls.PARSERS[file_ext]
        return parser_class(file_path)
    
    @classmethod
    def get_supported_formats(cls):
        """获取支持的文件格式列表"""
        return list(cls.PARSERS.keys())
    
    @classmethod
    def detect_format(cls, file_path):
        """自动检测文件格式"""
        # 基于文件扩展名
        file_ext = Path(file_path).suffix.lower()
        if file_ext in cls.PARSERS:
            return file_ext
        
        # 基于MIME类型
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            mime_to_ext = {
                'text/markdown': '.md',
                'text/plain': '.txt',
                'text/html': '.html',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
            }
            return mime_to_ext.get(mime_type, file_ext)
        
        return file_ext


# ==============================================
# XMind文件生成器（复用原有代码）
# ==============================================

def generate_content_xml(json_structure):
    """Generate content.xml content"""
    # 获取基本信息
    sheet_id = json_structure.get('id', 'default-sheet-id')
    sheet_title = json_structure.get('title', 'Sheet 1')
    root_topic = json_structure.get('rootTopic', {})
    
    # 生成XML头部
    xml_parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="no"?>',
        '<xmap-content xmlns="urn:xmind:xmap:xmlns:content:2.0" '
        'xmlns:fo="http://www.w3.org/1999/XSL/Format" '
        'xmlns:svg="http://www.w3.org/2000/svg" '
        'xmlns:xhtml="http://www.w3.org/1999/xhtml" '
        'xmlns:xlink="http://www.w3.org/1999/xlink" '
        'modified-by="Vana" timestamp="1503058545540" version="2.0">',
        f'<sheet id="{sheet_id}" modified-by="Vana" theme="0kdeemiijde6nuk97e4t0vpp54" timestamp="1503058545540">'
    ]
    
    # 生成根主题
    if root_topic:
        topic_id = root_topic.get('id', 'default-topic-id')
        topic_title = root_topic.get('title', 'Root Topic')
        
        xml_parts.append(
            f'<topic id="{topic_id}" modified-by="Vana" timestamp="1503058545484">'
            f'<title>{escape_xml_text(topic_title)}</title>'
        )
        
        # 递归生成子主题
        def generate_topics(topics, indent=0):
            if not topics:
                return []
            
            result = []
            result.append('<children><topics type="attached">')
            
            for topic in topics:
                topic_id = topic.get('id', f'topic-{hash(str(topic))}')
                title = topic.get('title', 'Topic')
                children = topic.get('children', {}).get('attached', [])
                
                result.append(f'<topic id="{topic_id}" modified-by="Vana" timestamp="1503058545484">')
                result.append(f'<title svg:width="500">{escape_xml_text(title)}</title>')
                
                if children:
                    result.extend(generate_topics(children, indent + 1))
                
                result.append('</topic>')
            
            result.append('</topics></children>')
            return result
        
        # 生成子主题
        attached_topics = root_topic.get('children', {}).get('attached', [])
        if attached_topics:
            xml_parts.extend(generate_topics(attached_topics))
        
        xml_parts.append('</topic>')
    
    # 添加扩展和标题
    xml_parts.extend([
        '<extensions><extension provider="org.xmind.ui.map.unbalanced">',
        '<content><right-number>-1</right-number></content>',
        '</extension></extensions>',
        f'<title>{escape_xml_text(sheet_title)}</title>',
        '</sheet>',
        '</xmap-content>'
    ])
    
    return '\n'.join(xml_parts)


def create_metadata():
    """Create metadata.json"""
    return {
        "dataStructureVersion": "2",
        "layoutEngineVersion": "3",
        "creator": {
            "name": "Vana",
            "version": "23.05.2004"
        }
    }


def create_manifest():
    """Create manifest.json"""
    return {
        "file-entries": {
            "content.json": {},
            "content.xml": {},
            "metadata.json": {},
            "Thumbnails/thumbnail.png": {}
        }
    }


def create_xmind_file(json_structure, output_file):
    """Create XMind file"""
    with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # 添加content.json - 注意：正常文件使用数组格式
        zip_file.writestr('content.json', json.dumps([json_structure], ensure_ascii=False, indent=2))
        
        # 添加content.xml
        content_xml = generate_content_xml(json_structure)
        zip_file.writestr('content.xml', content_xml)
        
        # 添加metadata.json
        zip_file.writestr('metadata.json', json.dumps(create_metadata(), ensure_ascii=False, indent=2))
        
        # 添加manifest.json
        zip_file.writestr('manifest.json', json.dumps(create_manifest(), ensure_ascii=False, indent=2))
        
        # 添加空缩略图目录
        zip_file.writestr('Thumbnails/', b'')
        
        # 添加默认缩略图（透明PNG）
        thumbnail_data = b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18\xd4c\x00\x00\x00\x00IEND\xaeB`\x82'
        zip_file.writestr('Thumbnails/thumbnail.png', thumbnail_data)


# ==============================================
# 兼容性函数（从markdown_to_xmind_converter.py合并）
# ==============================================

def parse_markdown_to_json(markdown_file):
    """兼容旧版本的Markdown解析函数"""
    parser = MarkdownParser(markdown_file)
    return parser.parse()

def main():
    """主函数 - 支持命令行参数"""
    import sys
    
    # 检查命令行参数
    if len(sys.argv) <= 1:
        print("[ERROR] 错误: 请提供输入文件路径")
        print("用法: python universal_xmind_converter.py <input_file>")
        print("支持的格式: .md, .txt, .html, .docx, .xlsx")
        return 1
    
    input_file = sys.argv[1]
    
    # 验证输入文件是否存在
    if not os.path.exists(input_file):
        print(f"[ERROR] 错误: 文件 '{input_file}' 不存在")
        return 1
    
    # 使用ParserFactory自动检测格式并转换
    try:
        factory = ParserFactory()
        parser = factory.get_parser(input_file)
        
        print(f"正在解析文件: {input_file}")
        json_structure = parser.parse()
        
        # 生成输出文件名（输出到output目录）
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_dir = "output"
        
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        output_file = os.path.join(output_dir, f"{base_name}.xmind")
        
        print(f"正在创建XMind文件: {output_file}")
        create_xmind_file(json_structure, output_file)
        
        print("[SUCCESS] 转换完成！")
        print(f"输出文件: {output_file}")
        
    except Exception as e:
        print(f"[ERROR] 转换失败: {str(e)}")
        return 1


if __name__ == '__main__':
    main()


# ==============================================
# 主函数
# ==============================================

def main():
    """Main function"""
    import sys
    
    # 检查命令行参数
    if len(sys.argv) < 2 or sys.argv[1] in ['-h', '--help', 'help']:
        print("=" * 60)
        print("Universal XMind Converter - 多格式思维导图转换器")
        print("=" * 60)
        print("\n支持的文件格式:")
        print("  • Markdown     (.md)     - 标题层级转换")
        print("  • 文本大纲     (.txt)    - 缩进格式大纲")
        print("  • HTML网页     (.html)   - 标题和列表结构")
        print("  • Word文档     (.docx)   - 标题样式转换")
        print("  • Excel表格    (.xlsx)   - 多列层级结构")
        print("\n使用方法:")
        print("  python universal_xmind_converter.py <输入文件> [输出文件]")
        print("\n示例:")
        print("  python universal_xmind_converter.py document.md")
        print("  python universal_xmind_converter.py outline.txt mymap.xmind")
        print("  python universal_xmind_converter.py data.xlsx")
        print("\n自动识别:")
        print("  无扩展名文件会自动检测格式")
        print("\n依赖安装:")
        print("  pip install beautifulsoup4 python-docx openpyxl")
        print("=" * 60)
        return 0
    
    if len(sys.argv) < 2:
        print("[ERROR] 错误: 请提供输入文件路径")
        print("用法: python universal_xmind_converter.py <input_file>")
        print("\n支持的文件格式:")
        for fmt in ParserFactory.get_supported_formats():
            print(f"  {fmt}")
        return 1
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 验证输入文件是否存在
    if not os.path.exists(input_file):
        print(f"[ERROR] 错误: 文件 '{input_file}' 不存在")
        return 1
    
    # 检测文件格式
    file_format = ParserFactory.detect_format(input_file)
    print(f"检测到文件格式: {file_format}")
    
    try:
        # 获取相应的解析器
        parser = ParserFactory.get_parser(input_file)
        print(f"使用解析器: {parser.__class__.__name__}")
        
        # 解析文件
        print(f"正在解析文件: {input_file}")
        json_structure = parser.parse()
        
        # 生成输出文件名
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_dir = "output"
        
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        output_file = os.path.join(output_dir, f"{base_name}.xmind")
        
        # 创建XMind文件
        print(f"正在创建XMind文件: {output_file}")
        create_xmind_file(json_structure, output_file)
        
        print("[SUCCESS] 转换完成！")
        print(f"📁 输出文件: {output_file}")
        
    except ImportError as e:
        print(f"[ERROR] 缺少依赖包: {e}")
        print("请安装相应的依赖包:")
        print("  pip install beautifulsoup4  # HTML解析")
        print("  pip install python-docx   # Word文档解析")
        print("  pip install openpyxl      # Excel解析")
        return 1
        
    except Exception as e:
        print(f"[ERROR] 转换失败: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    main()