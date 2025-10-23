from wowool.document.document_interface import DocumentInterface, MT_PLAINTEXT
import docx


def get_heading_markup_level(paragraph):
    if paragraph.style.style_id and paragraph.style.style_id.startswith("Heading"):
        return int(paragraph.style.style_id[7])
    return None


class DocxFileInputProvider(DocumentInterface):

    def __init__(self, fid, metadata: dict | None = None, add_heading: bool = True, **kwargs):
        self._id = str(fid)
        self._metadata = metadata if metadata is not None else {}
        self.add_heading = add_heading

    @property
    def id(self) -> str:
        return self._id

    @property
    def mime_type(self) -> str:
        return MT_PLAINTEXT

    @property
    def encoding(self) -> str:
        return "utf-8"

    @property
    def data(self):
        doc = docx.Document(self.id)
        text = ""
        for paragraph in doc.paragraphs:
            if self.add_heading:
                if markup_level := get_heading_markup_level(paragraph):
                    text += f"{'#' * markup_level} {paragraph.text}\n\n"
                else:
                    text += paragraph.text + "\n\n"
            else:
                text += paragraph.text + "\n\n"

        return text

    @property
    def metadata(self) -> dict:
        return self._metadata
