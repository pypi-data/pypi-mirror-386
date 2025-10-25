from pathlib import Path
import logging
from typing import Union

# Import the necessary third-party library
from docx import Document
from docx.opc.exceptions import OpcError

logger = logging.getLogger("ELESS.OfficeParser")


def parse_docx(file_path: Path) -> str:
    """
    Extracts all text content from a Microsoft Word (.docx) file.

    Args:
        file_path: The Path object pointing to the DOCX file.

    Returns:
        A single string containing the concatenated text from all paragraphs and tables,
        or an empty string if extraction fails.
    """
    text_content = []

    try:
        # Load the document
        document = Document(file_path)

        # 1. Extract text from paragraphs
        for paragraph in document.paragraphs:
            # Check if the paragraph is part of a list or header and include its text
            if paragraph.text:
                text_content.append(paragraph.text)

        # 2. Extract text from tables (important for structured data)
        for table in document.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                # Use a pipe | separator for cells and a newline for rows
                table_text.append(" | ".join(row_text))

            if table_text:
                text_content.append(
                    "\n\n---TABLE START---\n"
                    + "\n".join(table_text)
                    + "\n---TABLE END---\n"
                )

        # Join all extracted text with a paragraph separator
        return "\n\n".join(text_content).strip()

    except FileNotFoundError:
        logger.error(f"DOCX file not found at: {file_path}")
    except OpcError:
        logger.error(f"File {file_path.name} is corrupted or not a valid DOCX file.")
    except Exception as e:
        logger.error(
            f"An unexpected error occurred while parsing {file_path.name}: {e}"
        )

    return ""
