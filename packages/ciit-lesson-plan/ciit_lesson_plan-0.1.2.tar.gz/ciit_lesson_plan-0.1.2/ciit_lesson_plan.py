from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import re
import argparse

# 课前，课中等字体设置
def _color_white(merged_cell: _Cell):
  """
  设置合并单元格的文本居中、背景色为 #1A5F88，并设置字体颜色为白色且加粗。
  
  参数:
  merged_cell (_Cell): 要设置格式的表格单元格对象
  """
  # 设置单元格文本水平居中
  paragraph = merged_cell.paragraphs[0]
  paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  # 设置背景色为 #1A5F88
  tc = merged_cell._tc
  tcPr = tc.get_or_add_tcPr()
  shd = OxmlElement('w:shd')
  shd.set(qn('w:fill'), '1A5F88')  # 使用十六进制颜色码
  tcPr.append(shd)

  # 设置字体颜色为白色
  run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
  font = run.font
  font.color.rgb = RGBColor(255, 255, 255)
  run.font.size = Pt(16)  # 设置字体大小为16pt
  run.font.name = '微软雅黑'  # 设置字体为"微软雅黑"
  run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 确保中文字体兼容
  run.bold = True  # 添加这一行以实现加粗

def _read_markdown(file_path):
  with open(file_path, 'r', encoding='utf-8') as file:
    return file.read()

def _parse_markdown(content):
  data = {}
  current_key = None
  for line in content.splitlines():
    if line.startswith('# '):
      current_key = line[2:].strip()
      data[current_key] = ''
    elif line.startswith('## '):
      current_key = line[3:].strip()
      data[current_key] = ''
    elif line.startswith('### '):
      current_key = line[4:].strip()
      data[current_key] = ''  
    elif current_key:
      data[current_key] += line.strip() + '\n'
  return data

def convert_md_to_docx(template_path, markdown_path):
  template_path = 'template.docx'
  output_dir = '.'  # 输出目录
  document = Document(template_path)

  # 获取文档中的所有内容元素（段落和表格）
  all_elements = []
  for block in document.element.body:
    if block.tag == qn('w:p'):  # 是段落
      all_elements.append({'type': 'paragraph', 'element': block})
    elif block.tag == qn('w:tbl'):  # 是表格
      all_elements.append({'type': 'table', 'element': block})

  # 找到第17个段落的索引
  last_paragraph_to_keep_index = -1
  paragraph_count = 0
  for i, item in enumerate(all_elements):
    if item['type'] == 'paragraph':
      paragraph_count += 1
      if paragraph_count == 17:
        last_paragraph_to_keep_index = i
        break

  # 如果文档中的段落少于17个，则不删除任何段落
  if last_paragraph_to_keep_index == -1 and paragraph_count < 17:
    print(f"文档中的段落少于 17 个，不执行删除操作。")
  else:
    # 从第17个段落之后开始删除
    # 从后往前删除，避免索引问题
    for i in range(len(all_elements) - 1, last_paragraph_to_keep_index, -1):
      element_to_delete = all_elements[i]['element']
      parent = element_to_delete.getparent()
      parent.remove(element_to_delete)

  # 读取Markdown文件内容
  markdown_content = _read_markdown(markdown_path)
  parsed_data = _parse_markdown(markdown_content)

  # 提取项目名称
  project_name = parsed_data.get("项目名称", "").strip()

  # 从输入文件名中提取编号
  match = re.search(r'ja-(\d+)-out\.md', markdown_path)
  if match:
    number = match.group(1)
  else:
    number = "1"

  # 构建输出文件名
  output_path = f"{output_dir}/{number}-{project_name}-教案.docx"

  # 创建表格，第一行是2-5列合并，第二行是2-3列合并，第三行是3-5列合并
  table = document.add_table(rows=10, cols=5)
  table.style = 'Table Grid'
  table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  # 设置列宽,默认表格总宽度是20.
  col_widths = [1.0, 4.5, 7, 4, 3.5]
  for i, width_cm in enumerate(col_widths):
    table.columns[i].width = Cm(width_cm)
    for row in table.rows:
      row.cells[i].width = Cm(width_cm)

  # 合并单元格（保持原有逻辑）
  cell_0_0 = table.cell(0, 0)
  cell_0_1 = table.cell(0, 1)
  cell_0_1.merge(table.cell(0, 4))  # 第0行：1-4列合并

  cell_1_1 = table.cell(1, 1)
  cell_1_1.merge(table.cell(1, 2))  # 第1行：1-2列合并

  cell_2_2 = table.cell(2, 2)
  cell_2_2.merge(table.cell(2, 4))  # 第2行：2-4列合并

  cell_3_2 = table.cell(3, 2)
  cell_3_2.merge(table.cell(3, 4))  # 第3行：2-4列合并

  cell_4_2 = table.cell(4, 2)
  cell_4_2.merge(table.cell(4, 4))  # 第4行：2-4列合并

  cell_5_1 = table.cell(5, 1)
  cell_5_1.merge(table.cell(5, 4))  # 第5行：1-4列合并

  cell_6_1 = table.cell(6, 1)
  cell_6_1.merge(table.cell(6, 4))  # 第6行：1-4列合并

  cell_7_1 = table.cell(7, 1)
  cell_7_1.merge(table.cell(7, 4))  # 第7行：1-4列合并

  cell_8_1 = table.cell(8, 1)
  cell_8_1.merge(table.cell(8, 4))  # 第8行：1-4列合并

  cell_9_1 = table.cell(9, 1)
  cell_9_1.merge(table.cell(9, 4))  # 第9行：1-4列合并

  # 设置内容（从Markdown文件中读取）
  cell_0_0.text = "项目名称"
  cell_0_1.text = project_name

  cell_1_0 = table.cell(1, 0)
  cell_1_0.text = "授课类型"
  cell_1_1.text = parsed_data.get("授课类型", "").strip()

  cell_1_3 = table.cell(1, 3)
  cell_1_4 = table.cell(1, 4)
  cell_1_3.text = parsed_data.get("授课周次", "").strip()
  cell_1_4.text = parsed_data.get("授课学时", "").strip()

  cell_2_0 = table.cell(2, 0)
  cell_2_1 = table.cell(2, 1)
  cell_2_0.text = "教学目标"
  cell_2_1.text = "知识目标："
  cell_2_2.text = parsed_data.get("知识目标", "").strip()

  cell_3_0 = table.cell(3, 0)
  cell_3_0.merge(table.cell(2, 0))  # 合并到上一行
  cell_3_1 = table.cell(3, 1)
  cell_3_1.text = "能力目标："
  cell_3_2.text = parsed_data.get("能力目标", "").strip()

  cell_4_0 = table.cell(4, 0)
  cell_4_0.merge(table.cell(3, 0))  # 继续合并
  cell_4_1 = table.cell(4, 1)
  cell_4_1.text = "素质目标："
  cell_4_2.text = parsed_data.get("素质目标", "").strip()

  cell_5_0 = table.cell(5, 0)
  cell_5_0.text = "学情分析"
  cell_5_1.text = parsed_data.get("学情分析", "").strip()

  cell_6_0 = table.cell(6, 0)
  cell_6_0.text = "教学重点"
  cell_6_1.text = parsed_data.get("教学重点", "").strip()

  cell_7_0 = table.cell(7, 0)
  cell_7_0.text = "教学难点"
  cell_7_1.text = parsed_data.get("教学难点", "").strip()

  cell_8_0 = table.cell(8, 0)
  cell_8_0.text = "教学方法"
  cell_8_1.text = parsed_data.get("教学方法", "").strip()

  cell_9_0 = table.cell(9, 0)
  cell_9_0.text = "教材资源"
  cell_9_1.text = parsed_data.get("教材资源", "").strip()

  # 👇 设置外边框为 12
  tbl = table._element
  # 修正：使用正确的属性获取方法
  tblPr = tbl.tblPr
  if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)

  tblBorders = OxmlElement('w:tblBorders')

  for border_name in ['top', 'bottom', 'left', 'right']:
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '12')   
    border.set(qn('w:color'), 'auto')
    border.set(qn('w:space'), '0')
    tblBorders.append(border)

  tblPr.append(tblBorders)

  # 👇 统一设置字体格式 + 垂直居中
  for row_idx, row in enumerate(table.rows):
    for col in range(5):
      cell = row.cells[col]
      cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

      for paragraph in cell.paragraphs:
        for run in paragraph.runs:
          run.font.name = '宋体'
          run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
          run.font.size = Pt(10.5)
          run.bold = False

          if col == 0:
            run.font.size = Pt(12)
            run.bold = True

          if col == 1 and row_idx in [2, 3, 4]:
            run.font.size = Pt(12)
            run.bold = True

  # 设置所有行高
  for row in table.rows:
    row.height = Cm(1.5)

  # 添加表格，共5列，第一行内容为"课前"，5列合并为1列
  table = document.add_table(rows=18, cols=5)  # 修改：将行数改为8行
  table.style = 'Table Grid'
  table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

  # 设置列宽
  col_widths = [1.0, 14.0, 1.8, 1.2, 2]  # 每列宽度均为4厘米
  for i, width_cm in enumerate(col_widths):
    table.columns[i].width = Cm(width_cm)
    for row in table.rows:
      row.cells[i].width = Cm(width_cm)

  # 合并第一行的5列
  cell_0_0 = table.cell(0, 0)
  cell_0_4 = table.cell(0, 4)
  merged_cell = cell_0_0.merge(cell_0_4)
  merged_cell.text = "课前"

  _color_white(merged_cell)

  # 第二行内容为"教学内容"、"教学活动", "设计意图"
  # 合并第1-2列为"教学内容"
  cell_1_0 = table.cell(1, 0)
  cell_2_1 = table.cell(2, 1)
  cell_2_1.merge(cell_1_0)
  cell_1_0.text = "教学内容"

  # 合并第3-4列为"教学活动"
  cell_1_2 = table.cell(1, 2)
  cell_1_3 = table.cell(1, 3)
  cell_1_2.merge(cell_1_3)
  cell_1_2.text = "教学活动"

  cell_1_4 = table.cell(1, 4)
  cell_1_4.text = "设计意图"

  # 设置第二行的行高为0.6 cm
  table.rows[1].height = Cm(0.6)

  # 第三行设置表头内容
  cell_2_2 = table.cell(2, 2)
  cell_2_3 = table.cell(2, 3)
  cell_2_4 = table.cell(2, 4)
  cell_2_4.merge(cell_1_4)

  cell_2_2.text = "学生活动"
  cell_2_3.text = "教师活动"

  # 设置第三行的行高为1.9 cm
  table.rows[2].height = Cm(1.9)

  # 第四行：添加新内容
  cell_3_0 = table.cell(3, 0)
  cell_3_1 = table.cell(3, 1)
  cell_3_0.merge(cell_3_1)  # 合并第一列和第二列
  cell_3_0.text = parsed_data.get("课前:教学内容", "").strip()

  cell_3_2 = table.cell(3, 2)
  cell_3_2.text = parsed_data.get("课前:学生活动", "").strip()

  cell_3_3 = table.cell(3, 3)
  cell_3_3.text = parsed_data.get("课前:教师活动", "").strip()

  cell_3_4 = table.cell(3, 4)
  cell_3_4.text = parsed_data.get("课前:设计意图", "").strip()

  # 合并第5行的5列
  cell_4_0 = table.cell(4, 0)
  cell_4_4 = table.cell(4, 4)
  merged_cell = cell_4_0.merge(cell_4_4)
  merged_cell.text = "课中"

  _color_white(merged_cell)

  # 第6行内容为"教学内容"、"教学活动", "设计意图"
  # 合并第1-2列为"教学内容"
  table.cell(5,0).merge(table.cell(6,0))
  table.cell(5,0).text = "教学环节"

  table.cell(5,1).merge(table.cell(6,1))
  table.cell(5,1).text = "教学内容"

  # 合并第3-4列为"教学活动"
  cell_5_2 = table.cell(5, 2)
  cell_5_3 = table.cell(5, 3)
  cell_5_2.merge(cell_5_3)
  cell_5_2.text = "教学活动"

  table.cell(5,4).merge(table.cell(6,4))
  table.cell(5,4).text = "设计意图"

  # 第7行设置表头内容

  cell_6_2 = table.cell(6, 2)
  cell_6_3 = table.cell(6, 3)
  cell_6_4 = table.cell(6, 4)
  cell_6_4.merge(cell_6_4)

  cell_6_2.text = "学生活动"
  cell_6_3.text = "教师活动"

  # 项目导入
  table.cell(7,0).text = "项目导入"
  table.cell(7,1).text = parsed_data.get("项目导入:教学内容", "").strip()
  table.cell(7,2).text = parsed_data.get("项目导入:学生活动", "").strip()
  table.cell(7,3).text = parsed_data.get("项目导入:教师活动", "").strip()
  table.cell(7,4).text = parsed_data.get("项目导入:设计意图", "").strip()

  # 内容展开
  table.cell(8,0).text = "内容展开"
  table.cell(8,1).text = parsed_data.get("内容展开:教学内容", "").strip()
  table.cell(8,2).text = parsed_data.get("内容展开:学生活动", "").strip()
  table.cell(8,3).text = parsed_data.get("内容展开:教师活动", "").strip()
  table.cell(8,4).text = parsed_data.get("内容展开:设计意图", "").strip()

  # 课堂小结
  table.cell(9,0).text = "课堂小结"
  table.cell(9,1).text = parsed_data.get("课堂小结:教学内容", "").strip()
  table.cell(9,2).text = parsed_data.get("课堂小结:学生活动", "").strip()
  table.cell(9,3).text = parsed_data.get("课堂小结:教师活动", "").strip()
  table.cell(9,4).text = parsed_data.get("课堂小结:设计意图", "").strip()

  # 课后
  merged_cell = table.cell(10,0)
  merged_cell.merge(table.cell(10,4))
  merged_cell.text = "课后"

  _color_white(merged_cell)

  # 教学活动
  table.cell(11,0).merge(table.cell(12,1))
  table.cell(11,0).text = "教学内容"
  table.cell(11,2).merge(table.cell(11,3))
  table.cell(11,2).text = "教学活动"
  table.cell(11,4).merge(table.cell(12,4))
  table.cell(11,4).text = "设计意图"
  table.cell(12,2).text = "学生活动"
  table.cell(12,3).text = "教师活动"

  # 教学内容
  table.cell(13,0).merge(table.cell(13,1))
  table.cell(13,0).text = parsed_data.get("课后:教学内容", "").strip()
  table.cell(13,2).text = parsed_data.get("课后:学生活动", "").strip()
  table.cell(13,3).text = parsed_data.get("课后:教师活动", "").strip()
  table.cell(13,4).text = parsed_data.get("课后:设计意图", "").strip()

  # 教学反思
  merged_cell = table.cell(14, 0)
  merged_cell.merge(table.cell(14, 4))
  merged_cell.text = "教学反思"

  _color_white(merged_cell)

  # 教学效果
  table.cell(15,0).text ="教学效果"
  table.cell(15,1).merge(table.cell(15,4))
  table.cell(15,1).text = parsed_data.get("教学反思:教学效果", "").strip()

  # 诊断改进
  table.cell(16,0).merge(table.cell(17,0))
  table.cell(16,0).text = "诊断改进"
  table.cell(16,1).merge(table.cell(16,4))
  table.cell(16,1).text = parsed_data.get("教学反思:诊断", "").strip()
  table.cell(17,1).merge(table.cell(17,4))
  table.cell(17,1).text = parsed_data.get("教学反思:改进", "").strip()

  # 设置第二行的行高为2.5 cm
  table.rows[16].height = Cm(2)
  table.rows[17].height = Cm(2)

  # 设置表格字体样式和垂直居中
  tbl = table._element
  tblPr = tbl.tblPr
  if tblPr is None:
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)

  tblBorders = OxmlElement('w:tblBorders')

  for border_name in ['top', 'bottom', 'left', 'right']:
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '12')  # 0.5 磅 = 14/8 磅单位
    border.set(qn('w:color'), 'auto')
    border.set(qn('w:space'), '0')
    tblBorders.append(border)

  tblPr.append(tblBorders)

  # 设置表格字体样式和垂直居中（按单元格进行判断）
  special_box = ["课前", "课中", "课后", "教学反思"]
  title_box = ["教学内容", "教学活动", "学生活动", "教师活动", "设计意图", "教学环节",
         "项目导入", "内容展开", "课堂小结", "教学效果", "诊断改进"]

  for row in table.rows:
    for cell in row.cells:
      cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      cell_text = cell.text.strip()
      
      if cell_text in special_box:
        continue
      
      for paragraph in cell.paragraphs:
        for run in paragraph.runs:
          run.font.name = '宋体'
          run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
          
          if cell_text in title_box:
            run.font.size = Pt(12)
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
          else:
            run.font.size = Pt(10.5)
            run.bold = False
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

  # 添加空行
  paragraph = document.add_paragraph('')

  # 添加段落，右对齐
  paragraph = document.add_paragraph('制订时间: 2025 年 9 月')
  paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
  run = paragraph.runs[0]
  run.font.size = Pt(10.5)  # 字号16pt
  run.font.name = 'Calibri'
  run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

  # 保存文档
  document.save(output_path)

if __name__ == "__main__":

  parser = argparse.ArgumentParser(description="Convert Markdown to DOCX.")
  parser.add_argument("input_markdown_file", help="Path to the input Markdown file.")
  parser.add_argument("--template", default="template.docx", help="Path to the template DOCX file.")
  parser.add_argument("--output_dir", default=".", help="Output directory for the generated DOCX file.")
  args = parser.parse_args()

  convert_md_to_docx(args.template, args.input_markdown_file)
