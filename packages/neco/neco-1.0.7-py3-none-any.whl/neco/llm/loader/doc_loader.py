import base64
import tempfile

import docx
from langchain_core.documents import Document
from loguru import logger
from tqdm import tqdm

from neco.ocr.base_ocr import BaseOCR


class DocLoader:
    def __init__(self, file_path: str, ocr: BaseOCR, mode: str = 'full'):
        """
        mode: full, paragraph
        """

        self.file_path = file_path
        self.mode = mode
        self.ocr = ocr

    def load(self):
        logger.info(f"开始解析[{self.file_path}]的文档,解析模式为[{self.mode}]")

        try:
            document = docx.Document(self.file_path)
        except Exception as e:
            logger.error(f"加载文档[{self.file_path}]失败: {e}")
            return []

        paragraphs = document.paragraphs

        docs = []

        if self.mode == 'full':
            docs = self.full_mode_parser(paragraphs)

        elif self.mode == 'paragraph':
            docs = self.paragraph_mode_parse(paragraphs)

        else:
            logger.error(f"未知的解析模式: {self.mode}")
            return []

        tables = document.tables
        if tables:
            logger.info(f"检测到[{self.file_path}]中有[{len(tables)}]个表格,开始解析表格")
            for table in tqdm(tables, desc=f"解析[{self.file_path}]的表格"):
                docs.append(Document(self.table_to_md(table),
                                     metadata={"format": "table"}))

        # 提取图片并使用OCR识别
        if self.ocr is not None:
            logger.info(f"解析任务[{self.file_path}]启用了OCR识别,开始提取图片")
            for rel in document.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    try:
                        image_base64 = base64.b64encode(
                            image_data).decode('utf-8')
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=True) as temp_img:
                            temp_img.write(image_data)
                            temp_img.flush()
                            ocr_result = self.ocr.predict(temp_img.name)
                            docs.append(
                                Document(ocr_result,
                                         metadata={
                                             "format": "image", "image_base64": image_base64
                                         })
                            )
                    except Exception as e:
                        logger.error(f"处理图片失败: {e}")
                        continue

        return docs

    def table_to_md(self, table):
        # Converts a docx table to markdown format
        md_table = []
        for row in table.rows:
            md_row = '| ' + ' | '.join(cell.text for cell in row.cells) + ' |'
            md_table.append(md_row)
        return '\n'.join(md_table)

    def paragraph_mode_parse(self, paragraphs):
        current_doc = None
        docs = []

        for paragraph in tqdm(paragraphs, desc=f"解析[{self.file_path}]的段落"):
            if any(heading in paragraph.style.name for heading in ('Heading', '标题')):
                if current_doc is not None:
                    docs.append(Document(current_doc.strip()))
                current_doc = paragraph.text.strip() + "\n"  # Start a new
            else:
                if current_doc is not None:
                    current_doc += paragraph.text.strip() + "\n"
                else:
                    current_doc = paragraph.text.strip() + "\n"

        if current_doc:
            docs.append(Document(current_doc.strip()))

        return docs

    def full_mode_parser(self, paragraphs):
        docs = []

        full_text = ""
        for paragraph in tqdm(paragraphs, desc=f"解析[{self.file_path}]的段落"):
            full_text += paragraph.text

        docs.append(Document(full_text))
        return docs
